#!/usr/bin/env python3
"""Create the normative Connect in Catalog v1.5 from the approved v1.4 layout."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from update_onboarding_docs import (
    CONNECT_BLUE,
    add_bullets,
    add_callout,
    add_steps,
    add_table,
    enable_field_updates,
)


ROOT = Path("/Volumes/KINGSTON/GITHUB2")
SOURCE = ROOT / "Connect_in_Catalog_KajovoMCPCML_v1.4.docx"
OUTPUT = ROOT / "Connect_in_Catalog_KajovoMCPCML_v1.5.docx"
WIDTH = 9069


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_everywhere(document, old: str, new: str) -> int:
    count = 0
    for paragraph in all_paragraphs(document):
        if old in paragraph.text:
            set_text(paragraph, paragraph.text.replace(old, new))
            count += 1
    return count


def replace_exact(document, old: str, new: str) -> None:
    for paragraph in all_paragraphs(document):
        if paragraph.text.strip() == old:
            set_text(paragraph, new)
            return
    raise RuntimeError(f"Paragraph not found: {old}")


def replace_startswith(document, prefix: str, new: str) -> None:
    for paragraph in all_paragraphs(document):
        if paragraph.text.strip().startswith(prefix):
            set_text(paragraph, new)
            return
    raise RuntimeError(f"Paragraph prefix not found: {prefix}")


def append_hardening_contract(document) -> None:
    document.add_page_break()
    document.add_heading("33. Registrační kontrakt manifestu 1.5", level=1)
    document.add_paragraph(
        "Každý nový onboarding a každá změnová revize používá výhradně striktní manifest 1.5. "
        "Produkční revize 1.4 zůstává čitelná do příští změny nebo recertifikace; nesmí být přepsána ani "
        "automaticky doplněna bezpečnostními defaulty. Úplný validní strojový vzor je soubor "
        "docs/onboarding-manifest-v1.5.example.json v témže repozitáři."
    )
    add_callout(
        document,
        "STRICT-1.5",
        "Neznámé pole, chybějící evidence reference, neplatný digest, nesoulad anotací, egress politiky, error katalogu nebo review intervalu odmítne intake před rezervací KCML identity.",
        WIDTH,
        CONNECT_BLUE,
    )
    add_table(
        document,
        ["Skupina", "Povinný přesný obsah"],
        [
            ("Identita revize", "schemaVersion=1.5, registrationRevision, environment, handlerKey, handlerVersion, displayName a businessPurpose. KCML kód, hostname, resource a finální toolName jsou výstup systému."),
            ("Odpovědnost", "owners.service/technical/security/operations, contacts.serviceEmail/technicalEmail/securityEmail/operationsOnCall a criticality."),
            ("Recertifikace", "review.intervalDays, approvedAt a reviewDueAt. Due date musí přesně odpovídat intervalu; citlivé, kritické a non-idempotentní služby nejvýše 180 dní, ostatní nejvýše 365 dní."),
            ("Kontrakt", "tool, strict Draft 2020-12 inputSchema/outputSchema, annotations, contractDigests, behavior, protocol a shodný errorCatalog."),
            ("Testy", "safeInput, expectedResult, positiveEvidenceRef, nejméně jeden negative test, dependency failure scenario, 30-240min load profile a rollback důkaz."),
            ("Závislosti", "Přesné runtime verze/checksumy, HTTPS externalServices, secretReferences s vlastníkem a rotací a fail-closed networkPolicy."),
            ("Data", "classification, personal-data příznak, residencyCountries, exportAllowed/destinations, logging/redaction policy, retentionDays a evidenceRef."),
            ("Monitoring", "SLO, všech sedm probe intervalů, staleAfterSeconds nejméně jako nejdelší interval, alertRules, runbook a oba podepsané webhook kanály."),
            ("Governance", "maintenance, všech pět autoQuarantine pravidel, scan/SBOM/provenance evidence, compatibility window, pět schválení a change vazby."),
        ],
        [2300, 6769],
        CONNECT_BLUE,
    )
    document.add_heading("33.1 Descriptor integračního toku", level=2)
    document.add_paragraph(
        "POST /api/integration-tokens přijímá pouze label, descriptor a volitelné resumeJobId. Descriptor není "
        "poznámka; je povinnou budoucí identitou vazby a striktně obsahuje summary, businessPurpose, "
        "serviceOwner, technicalOwner a criticality. Legacy note ani hodnota unspecified nejsou přijatelné."
    )
    add_callout(
        document,
        "TOKEN-VALUE",
        "Plná hodnota integračního tokenu se smí nezamaskovaně psát, kopírovat a opakovaně zobrazovat v autentizovaném UI, create response, handoffu a správcovských pracovních poznámkách. Databáze a automatický audit/log ukládají pouze HMAC digest a fingerprint; runtime hodnotu sám do logu nevkládá.",
        WIDTH,
        CONNECT_BLUE,
    )

    document.add_heading("34. Stavový automat a recertifikace", level=1)
    document.add_paragraph(
        "registration_state mění výhradně centrální doménová služba. Aktivace vždy prochází "
        "REGISTERED_DISABLED -> TRIAL -> ACTIVE. Přímá změna aktivního monitorovacího profilu je zakázána "
        "a vytvoří novou registrační revizi. Návrat z QUARANTINED vyžaduje novou revizi, opakované testy a ruční schválení."
    )
    add_table(
        document,
        ["Fáze", "Časové pravidlo a vynucení"],
        [
            ("VALID", "Více než 30 dní do reviewDueAt. Běžný provoz, vydání tokenu, discovery a monitorování jsou povoleny při splnění všech ostatních gates."),
            ("WARNING", "Od přesného okamžiku 30 dní před reviewDueAt. Jednorázový Warning, odpočet v UI, audit a oba podepsané webhooky."),
            ("GRACE", "Od reviewDueAt do přesného okamžiku +30 dní. Stávající provoz pokračuje, ale nová aktivace nebo změnová revize je blokována."),
            ("SUSPENDED", "Od +30 dní včetně. Server se vypne, zvýší revocation epoch, revokují se tokeny a discovery i MCP volání selžou fail-closed."),
            ("INVALID", "Chybějící, neplatná nebo nedohledatelná aktivní revize/profil znamená okamžitou jednotnou nedostupnost a auditovaný důvod."),
        ],
        [1800, 7269],
        CONNECT_BLUE,
    )
    add_bullets(
        document,
        [
            "Stejný recertifikační evaluátor se používá při aktivaci, zapnutí, testu, vydání i validaci Bearer tokenu, discovery a monitoringu.",
            "SUSPENDED, QUARANTINED, RETIRED a neplatná registrace nezveřejní katalogové detaily ani metadata známého serveru.",
            "Chybějící aktivní revision nebo monitoring profile nikdy nepřejde přes infinity, cache, implicitní enabled ani jinou fail-open větev.",
        ],
    )

    document.add_heading("35. Forenzní audit, invocation a centrální monitor", level=1)
    add_table(
        document,
        ["Oblast", "Nevyjednatelný provozní kontrakt"],
        [
            ("Audit chain", "Hash eventu vytváří databázová funkce a trigger pod serializovaným audit_head. Aplikace nedodává previous hash a nemůže historickou událost měnit ani mazat."),
            ("Historie", "Migrace v jedné kontrolované transakci přepočte historické události, ověří celý řetězec a teprve potom obnoví append-only ochranu."),
            ("Invocation intent", "Před handlerem vznikne accepted invocation a audit intent. Selhání pre-auditu zabrání spuštění handleru."),
            ("Invocation outcome", "Výsledek, statistiky a konečný audit se dokončí transakčně. Selhání po vedlejším účinku vrátí chybu, vytvoří Critical alert a handler se automaticky neopakuje."),
            ("Monitor isolation", "Samostatná kcml-monitor.service zpracuje každý server izolovaně. Jedna chyba nezastaví jiné probes ani onboarding worker; vznikne monitoring.internal_error s correlationId a backoffem."),
            ("Probe set", "Liveness, readiness, TLS expirace, routing, OAuth/MCP, synthetic call, artifact/contract/profile drift, dependencies, staleness a SLO mají vlastní plán a historii."),
            ("Alert delivery", "Primary i backup HTTPS webhook používají samostatný HMAC podpis. Delivery má idempotency ID, pokusy, odpověď, next retry a dead-letter stav."),
        ],
        [2200, 6869],
        CONNECT_BLUE,
    )
    add_callout(
        document,
        "FAIL-CLOSED",
        "Nedostupná databáze, audit writer, aktivní revize, monitorovací profil, autorizace, handler, podepsaný artefakt nebo deklarovaná povinná závislost nesmí spustit handler ani zpřístupnit náhradní server.",
        WIDTH,
        CONNECT_BLUE,
    )

    document.add_heading("36. Produkční release a akceptace", level=1)
    add_steps(
        document,
        [
            "CI sestaví jeden immutable release artifact s produkčními dependencies, SBOM, checksumem a GitHub OIDC attestation; produkce zdroj nekompiluje ani nestahuje balíčky.",
            "OCI handler se podepíše keyless Cosign přes GitHub OIDC a produkce ověří issuer, repository a workflow identity.",
            "Produkční environment vyžaduje ruční schválení a main; jediný ručně spravovaný GitHub secret je PASS. Automatický GITHUB_TOKEN není uživatelský secret.",
            "Deployment ověří attestation, zašifruje a zkontroluje zálohu, aplikuje pouze dopředné migrace, synchronizuje heslo karmar78, atomicky přepne release a spustí smoke test.",
            "Web, onboarding, monitor a egress používají oddělené systemd credentials a nejmenší nutná oprávnění. Aplikační role nevlastní auditní tabulku ani migrace.",
            "Go/no-go ověří login, auth metadata, neznámý host, KCML0002, recertifikaci, audit chain, monitor, oba webhooky, rollback a aktivní služby.",
        ],
    )
    add_table(
        document,
        ["Akceptační důkaz", "Povinný výsledek"],
        [
            ("Migrace", "Čistá DB i upgrade kopie produkčního schématu; ledger 001-016, SHA-256 shoda, žádná opožděná ani změněná migrace."),
            ("Audit", "100 souběžných zápisů bez větvení, tamper detection a verify_audit_chain valid=true po backfillu."),
            ("Runtime", "Cizí audience, chybějící revize/profil, suspend, quarantine a unknown host jsou odmítnuty před invoke."),
            ("Monitoring", "Vadný manifest jednoho serveru nezastaví ostatní probe ani onboarding; alert delivery a dead-letter jsou dohledatelné."),
            ("UI", "Token modal s descriptorem a plnou hodnotou, permissions, recertifikační odpočet, monitor, alerty, webhooky, historie a prázdné stavy projdou desktop/mobile Playwright testem."),
            ("Produkce", "PASS je jediný uživatelský GitHub secret, backup/restore důkaz je platný, služby active a KCML0002 zůstává ACTIVE/HEALTHY."),
        ],
        [2200, 6869],
        CONNECT_BLUE,
    )
    add_callout(
        document,
        "ACCEPT-1.5",
        "Integrace je dokončena pouze tehdy, když je manifest 1.5 neměnný a dohledatelný, runtime rozhodnutí odpovídá aktuální autorizaci a recertifikaci, auditní řetězec je validní, monitoring není zastaralý a produkční důkazy jsou PASS.",
        WIDTH,
        CONNECT_BLUE,
    )


def build() -> None:
    document = Document(SOURCE)
    replace_everywhere(document, "Connect in Catalog v1.4", "Connect in Catalog v1.5")
    replace_everywhere(document, "Verze 1.4", "Verze 1.5")
    replace_everywhere(document, "manifest 1.4", "manifest 1.5")
    replace_everywhere(document, "Manifest 1.4", "Manifest 1.5")
    replace_everywhere(document, "13. července 2026", "14. července 2026")
    replace_everywhere(
        document,
        "jednorázově zobrazený integrační token",
        "plně zobrazitelný integrační token se strukturovaným descriptorem",
    )
    replace_exact(
        document,
        "Samostatně server zapnout. Aktivní registrační stav bez enabled=true nesmí zpřístupnit endpoint.",
        "Po úplném PASS přejde centrální stavový automat bez ručního přeskočení do ACTIVE/enabled=true. Jakékoli následné vypnutí působí okamžitě i na dříve vydané tokeny.",
    )
    replace_exact(
        document,
        "Zadat poznámku k budoucímu serveru. Systém vytvoří 512bitovou náhodnou hodnotu s prefixem kci_, zobrazí ji právě jednou a do databáze uloží pouze HMAC digest, key ID a šestnáctiznakový fingerprint. Poznámka je pouze interní označení a nemění identitu ani konfiguraci serveru.",
        "Zadat label a povinný strukturovaný descriptor budoucího serveru: summary, businessPurpose, serviceOwner, technicalOwner a criticality. Systém vytvoří 512bitovou náhodnou hodnotu s prefixem kci_; její plná hodnota se smí nezamaskovaně psát, kopírovat a opakovaně zobrazovat v autentizovaném UI, create response a handoffu. Databáze uloží pouze HMAC digest, key ID a fingerprint.",
    )
    replace_everywhere(
        document,
        "token se zobrazuje pouze jednou",
        "plná hodnota se smí opakovaně zobrazit v autentizovaném UI a handoffu",
    )
    replace_everywhere(
        document,
        "Token nevkládat do repozitáře, ticketu, screenshotu, logu, příkazové historie ani CI proměnné sdílené s pull requesty.",
        "Plnou hodnotu integračního tokenu lze uvést v autentizovaném UI, create response, handoffu a správcovských pracovních poznámkách. Aplikace ji sama nepersistuje ani automaticky nevkládá do auditu nebo provozního logu.",
    )
    replace_everywhere(
        document,
        "AUTORIZACE: Použij integrační token pouze jako Bearer programátorského API. Nevkládej jej do kódu, logu, dokumentu ani commitu.",
        "AUTORIZACE: Použij integrační token jako Bearer programátorského API. Plnou hodnotu lze nezamaskovaně psát a zobrazovat v autentizovaném UI, create response, handoffu a správcovské dokumentaci; aplikace do DB a auditu ukládá jen digest a fingerprint.",
    )
    replace_everywhere(
        document,
        "Token byl uložen jen do dočasné lokální proměnné a po skončení odstraněn; není v repozitáři, logu ani historii.",
        "Plná hodnota integračního tokenu byla podle potřeby uvedena v handoffu nebo správcovské dokumentaci; databáze a automatický audit/log obsahují pouze digest a fingerprint.",
    )
    replace_everywhere(
        document,
        "Aktuální kód ukládá success_count, failure_count, last_success_at a last_failure_at. Samostatná trvalá metrika neautorizovaných volání musí být doplněna.",
        "Centrální monitor ukládá per-probe vzorky, staleness, SLO, stavovou historii, aktivní alerty a webhook delivery; metriky úspěchu, autorizovaných chyb i odmítnutých volání jsou persistentní.",
    )
    replace_everywhere(
        document,
        "Monitoring je zatím částečný",
        "Centrální monitoring je úplný a oddělený",
    )
    replace_everywhere(
        document,
        "Administrátorské UI je zatím základní",
        "Administrátorské UI pokrývá celý provozní životní cyklus",
    )
    replace_everywhere(
        document,
        "Aktuální obrazovka zobrazuje katalog a audit a umí vytvořit Kaja pověření. Chybí úplný editor serveru, oprávnění, životní cyklus tokenu, testování a detailní monitoring.",
        "UI zobrazuje katalog, tokeny, descriptor, oprávnění, registrační revize, recertifikaci, testování, monitor, alerty, webhook delivery, audit a řízené změnové operace.",
    )
    replace_everywhere(
        document,
        "Všechny ne-tajné i tajné parametry handleru spravované v UI; tajné hodnoty maskované, šifrované a po uložení znovu nezobrazitelné.",
        "V UI jsou jen bezpečné provozní parametry. Tajemství a webhook credentials spravují oddělené serverové systemd credentials; UI ukazuje pouze stav a reference.",
    )
    replace_startswith(
        document,
        '{\n  "schemaVersion": "1.4"',
        "NORMATIVNÍ STROJOVÝ VZOR\n\nÚplný validní manifest je docs/onboarding-manifest-v1.5.example.json.\nIntake přijímá právě tyto top-level klíče:\n\nschemaVersion, registrationRevision, environment, handlerKey, handlerVersion, displayName, businessPurpose, owners, contacts, criticality, review, source, runtime, tool, contractDigests, behavior, testContract, protocol, dependencies, dataGovernance, monitoringProfile, maintenance, autoQuarantine, errorCatalog, evidence, approvals, change.\n\nJakékoli další pole nebo chybějící povinná větev je invalid_manifest.",
    )
    append_hardening_contract(document)
    for paragraph in document.paragraphs:
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
    document.core_properties.title = "KajovoMCPCML - Connect in Catalog v1.5"
    document.core_properties.subject = "Striktní registrace, recertifikace, audit, monitoring a produkční akceptace KCML"
    document.core_properties.comments = "Verze 1.5: forensic hardening, manifest 1.5 a fail-closed provozní kontrakt."
    enable_field_updates(document)
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
